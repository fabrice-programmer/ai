"""Document service — handles file upload, validation, text extraction, and persistence.

Flow:
  Upload → Validate file (type, size) → Extract text → Save DB record → Store file safely
"""

import logging
import os
import uuid
from pathlib import Path
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename
from backend.extensions import db
from backend.models.document import Document
from backend.utils.error_handlers import AppError

logger = logging.getLogger(__name__)


# ── Public API ────────────────────────────────────────────────────


def upload_document(
    user_id: int,
    file: FileStorage,
    upload_folder: str,
    allowed_extensions: set | None = None,
    max_size_mb: int = 10,
    subdir: str = "documents",
) -> Document:
    """Full upload pipeline: validate → extract → persist → store.

    Parameters
    ----------
    user_id : int
        ID of the uploading user.
    file : FileStorage
        The uploaded file from the request.
    upload_folder : str
        Root upload directory from config.
    allowed_extensions : set | None
        Set of permitted extensions (e.g. {".docx"}).
    max_size_mb : int
        Maximum allowed file size in megabytes.
    subdir : str
        Subdirectory within *upload_folder* for storing files.

    Returns
    -------
    Document
        The newly created database record.

    Raises
    ------
    AppError
        If validation or storage fails with a user-friendly message.
    """
    # ── 1. Validate file ─────────────────────────────────────────
    validate_file(file, allowed_extensions or {".docx"}, max_size_mb)

    # ── 2. Extract text ──────────────────────────────────────────
    text = extract_text_from_file(file)

    # ── 3. Create database record (before saving file so we have doc.id) ─
    original_name = secure_filename(file.filename or "untitled.docx")
    doc = Document(
        user_id=user_id,
        filename=original_name,
        extracted_text=text,
        text_length=len(text) if text else 0,
        status="completed",
    )
    db.session.add(doc)
    db.session.flush()  # assign doc.id without committing yet

    # ── 4. Store file safely ─────────────────────────────────────
    stored_path = _store_file(file, upload_folder, subdir, doc.id, original_name)
    doc.stored_path = stored_path

    db.session.commit()

    logger.info(
        "Document %d ('%s', %d chars) uploaded by user %d -> %s",
        doc.id, original_name, doc.text_length, user_id, stored_path,
    )
    return doc


def get_document(doc_id: int) -> Document:
    """Retrieve a document by ID."""
    doc = db.session.get(Document, doc_id)
    if not doc:
        raise AppError("Document not found", 404)
    return doc


def get_user_documents(user_id: int) -> list[Document]:
    """List all documents belonging to a user, newest first."""
    return (
        Document.query.filter_by(user_id=user_id)
        .order_by(Document.upload_date.desc())
        .all()
    )


def delete_document(doc_id: int) -> None:
    """Delete a document record and its stored file.

    Raises AppError if the document does not exist.
    """
    doc = get_document(doc_id)

    # Remove the file from disk if it exists
    if doc.stored_path and os.path.isfile(doc.stored_path):
        try:
            os.remove(doc.stored_path)
            logger.info("Deleted file %s for document %d", doc.stored_path, doc_id)
        except OSError as exc:
            logger.warning("Could not delete file %s: %s", doc.stored_path, exc)

    db.session.delete(doc)
    db.session.commit()
    logger.info("Document %d deleted", doc_id)


# ── Validation ────────────────────────────────────────────────────


def validate_file(
    file: FileStorage,
    allowed_extensions: set[str],
    max_size_mb: int,
) -> None:
    """Validate file presence, type, and size.

    Parameters
    ----------
    file : FileStorage
        The uploaded file object.
    allowed_extensions : set of str
        Permitted file extensions (lowercase, with dot, e.g. {".docx"}).
    max_size_mb : int
        Maximum file size in megabytes.

    Raises
    ------
    AppError
        With appropriate HTTP status code and descriptive message.
    """
    if file is None:
        raise AppError("No file was provided in the upload request.", 400)

    # FileStorage.__bool__ returns False when filename is empty
    if not file or not file.filename:
        raise AppError("The uploaded file must have a valid filename.", 400)

    filename = file.filename
    if not filename.strip():
        raise AppError("The uploaded file must have a valid filename.", 400)

    ext = _get_extension(filename)
    if ext not in allowed_extensions:
        allowed_str = ", ".join(sorted(allowed_extensions))
        raise AppError(
            f"File type '{ext}' is not supported. "
            f"Only the following file types are allowed: {allowed_str}.",
            400,
        )

    # Check file size by seeking to end
    file.seek(0, os.SEEK_END)
    size_bytes = file.tell()
    file.seek(0)  # rewind for subsequent reads

    max_bytes = max_size_mb * 1024 * 1024
    if size_bytes > max_bytes:
        raise AppError(
            f"File exceeds the maximum allowed size of {max_size_mb} MB "
            f"(uploaded: {size_bytes / (1024 * 1024):.2f} MB).",
            413,  # Payload Too Large
        )


def extract_text_from_file(file: FileStorage) -> str:
    """Extract text content from an uploaded .docx file.

    Parameters
    ----------
    file : FileStorage
        The uploaded file (position reset to start before extraction).

    Returns
    -------
    str
        Extracted text content.

    Raises
    ------
    AppError
        If text extraction fails.
    """
    file.seek(0)

    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(file)
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)
    except ImportError:
        logger.exception("python-docx is not installed")
        raise AppError(
            "Document processing is unavailable. "
            "Please contact the administrator (python-docx missing).",
            500,
        )
    except Exception as exc:
        logger.exception("Failed to extract text from .docx file")
        raise AppError(
            f"Could not read the .docx file: {exc}", 400,
        )


# ── Internal helpers ──────────────────────────────────────────────


def _get_extension(filename: str) -> str:
    _, ext = os.path.splitext(filename)
    return ext.lower()


def _store_file(
    file: FileStorage,
    upload_folder: str,
    subdir: str,
    doc_id: int,
    original_name: str,
) -> str:
    """Persist the uploaded file to disk with a UUID-based name.

    Directory structure::

        <upload_folder>/<subdir>/<uuid>_<safe_name>

    Parameters
    ----------
    file : FileStorage
        The uploaded file (position reset to start).
    upload_folder : str
        Root upload directory.
    subdir : str
        Subdirectory within *upload_folder*.
    doc_id : int
        Document ID used in the stored filename for traceability.
    original_name : str
        Already-sanitised original filename.

    Returns
    -------
    str
        The absolute path to the stored file.

    Raises
    ------
    AppError
        If the target directory cannot be created or the file cannot be saved.
    """
    target_dir = Path(upload_folder) / subdir
    try:
        target_dir.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        logger.exception("Cannot create upload directory %s", target_dir)
        raise AppError(
            "The server could not prepare storage for the uploaded file. "
            "Please try again later.",
            500,
        ) from exc

    # Generate a unique filename to prevent collisions and path traversal
    unique_id = uuid.uuid4().hex[:12]
    safe_name = f"{doc_id}_{unique_id}_{original_name}"
    dest_path = target_dir / safe_name

    file.seek(0)
    try:
        file.save(str(dest_path))
    except OSError as exc:
        logger.exception("Failed to save file to %s", dest_path)
        raise AppError(
            "The uploaded file could not be saved due to a server error. "
            "Please try again later.",
            500,
        ) from exc

    return str(dest_path.resolve())