"""DOCX generation service — converts processed text into a .docx file.

Preserves:
  - Headings (detected via Markdown-style `#`/`##`/`###` prefixes or leading/trailing blank-lines pattern)
  - Paragraphs (regular text blocks)
  - References (a trailing reference/sources section)
"""

import logging
import re
import tempfile
from pathlib import Path
from typing import IO

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

logger = logging.getLogger(__name__)

# ── Default styling constants ──────────────────────────────────────

HEADING_STYLES: dict[int, dict] = {
    1: {"font_size": 18, "bold": True, "color": RGBColor(0x1A, 0x1A, 0x2E), "space_before": 24},
    2: {"font_size": 15, "bold": True, "color": RGBColor(0x2C, 0x2C, 0x54), "space_before": 18},
    3: {"font_size": 13, "bold": True, "color": RGBColor(0x3A, 0x3A, 0x6E), "space_before": 12},
}

BODY_FONT_SIZE = 11
BODY_FONT_NAME = "Calibri"
HEADING_FONT_NAME = "Calibri"

# Matches markdown-style headings: # Heading, ## Heading, ### Heading
HEADING_MD_RE = re.compile(r"^(#{1,3})\s+(.+)$", re.MULTILINE)

# ── Public API ─────────────────────────────────────────────────────


def generate_docx(
    text: str,
    output_path: str | Path | None = None,
    title: str | None = None,
) -> bytes | str:
    """Generate a .docx file from processed *text*.

    Parameters
    ----------
    text : str
        The processed text to convert.  Headings can be denoted with
        Markdown-style ``# `` / ``## `` / ``### `` prefixes, or detected
        heuristically (short lines surrounded by blank lines).
    output_path : str or Path, optional
        If provided, the .docx is saved to this path and the path is
        returned.  If omitted, the raw bytes are returned.
    title : str, optional
        Optional document title (rendered as a centred heading at the top).

    Returns
    -------
    bytes or str
        The .docx file as raw bytes when *output_path* is ``None``,
        otherwise the absolute path to the saved file.

    Raises
    ------
    OSError
        If the file cannot be written.
    """
    doc = DocxDocument()

    # ── Default font ──────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = BODY_FONT_NAME
    font.size = Pt(BODY_FONT_SIZE)

    # ── Optional title ────────────────────────────────────────────
    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run.font.name = HEADING_FONT_NAME
        # Small spacer after title
        doc.add_paragraph()

    # ── Parse and write content ───────────────────────────────────
    blocks = _parse_blocks(text)

    for block in blocks:
        _write_block(doc, block)

    # ── Return or save ────────────────────────────────────────────
    if output_path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info("DOCX saved to %s", output_path.resolve())
        return str(output_path.resolve())

    # Return raw bytes
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    doc.save(tmp_path)
    with open(tmp_path, "rb") as fh:
        raw = fh.read()
    Path(tmp_path).unlink(missing_ok=True)
    return raw


def generate_docx_from_blocks(
    blocks: list[dict],
    output_path: str | Path | None = None,
    title: str | None = None,
) -> bytes | str:
    """Generate a .docx file from a pre-parsed list of *blocks*.

    Each *block* is a dict with keys:

    - ``type`` : ``"heading1"`` | ``"heading2"`` | ``"heading3"`` | ``"paragraph"``
    - ``text`` : the text content

    This is the convenience entry-point when the caller has already
    structured the content (e.g. from an AI pipeline).

    Parameters, return value and exceptions are identical to
    :func:`generate_docx`.
    """
    doc = DocxDocument()

    style = doc.styles["Normal"]
    font = style.font
    font.name = BODY_FONT_NAME
    font.size = Pt(BODY_FONT_SIZE)

    if title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run.font.name = HEADING_FONT_NAME
        doc.add_paragraph()

    for block in blocks:
        _write_block(doc, block)

    if output_path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info("DOCX saved to %s", output_path.resolve())
        return str(output_path.resolve())

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    doc.save(tmp_path)
    with open(tmp_path, "rb") as fh:
        raw = fh.read()
    Path(tmp_path).unlink(missing_ok=True)
    return raw


# ── Block parsing ──────────────────────────────────────────────────


def _parse_blocks(text: str) -> list[dict]:
    """Split *text* into a list of structured blocks.

    Detection order:
      1. Markdown-style headings (``# Title``)
      2. Reference sections (lines containing ``references``, ``sources``, ``bibliography``)
      3. Everything else becomes a paragraph block.

    Returns
    -------
    list[dict]
        Each dict has ``type`` and ``text`` keys.
    """
    lines = text.split("\n")
    blocks: list[dict] = []
    buffer: list[str] = []

    def flush_buffer() -> None:
        """Dump the current paragraph buffer as a paragraph block."""
        nonlocal buffer
        content = "\n".join(buffer).strip()
        if content:
            blocks.append({"type": "paragraph", "text": content})
        buffer = []

    in_references = False
    reference_buffer: list[str] = []

    for line in lines:
        # ── Markdown heading detection ────────────────────────────
        md_match = HEADING_MD_RE.match(line)
        if md_match:
            flush_buffer()
            level = len(md_match.group(1))  # 1, 2, or 3
            heading_text = md_match.group(2).strip()
            blocks.append({"type": f"heading{level}", "text": heading_text})
            continue

        stripped = line.strip()

        # ── Reference section detection ───────────────────────────
        # If we see a line that looks like a reference heading, switch to
        # reference mode.  Everything below it is collected as one block.
        if (
            _is_reference_heading(stripped)
            or (in_references and _is_reference_line(stripped))
        ):
            if not in_references:
                flush_buffer()
                in_references = True
            # If this is the heading line itself, add it
            if _is_reference_heading(stripped):
                reference_buffer.append(stripped)
            else:
                if stripped:
                    reference_buffer.append(stripped)
            continue
        else:
            if in_references:
                # Flush collected references
                if reference_buffer:
                    blocks.append({"type": "paragraph", "text": "\n".join(reference_buffer)})
                    reference_buffer = []
                in_references = False

        # ── Normal paragraph lines ────────────────────────────────
        if stripped == "" and buffer:
            # Blank line → paragraph boundary
            flush_buffer()
        else:
            buffer.append(line)

    # Flush any remaining content
    if in_references and reference_buffer:
        blocks.append({"type": "paragraph", "text": "\n".join(reference_buffer)})
    else:
        flush_buffer()

    return blocks


def _is_reference_heading(stripped: str) -> bool:
    """Check if a stripped line is a reference/sources/bibliography heading."""
    lower = stripped.lower().rstrip(":")
    return lower in (
        "references",
        "sources",
        "works cited",
        "bibliography",
        "further reading",
    )


def _is_reference_line(stripped: str) -> bool:
    """Heuristic: a line that looks like a citation/reference entry."""
    if not stripped:
        return False
    # Common reference patterns: numbered "[1]", "(Author, 2020)",
    # starting with "Author. (2020)", or URLs
    if re.match(r"^\[\d+\]", stripped):
        return True
    if re.match(r"^[A-Z][a-z]+.+\(\d{4}", stripped):
        return True
    if "http" in stripped or "doi." in stripped:
        return True
    return False


# ── Block writing ──────────────────────────────────────────────────


def _write_block(doc: DocxDocument, block: dict) -> None:
    """Write a single structured block into the *doc*."""
    block_type = block.get("type", "paragraph")
    text = block.get("text", "")

    if not text.strip():
        return

    # ── Headings ──────────────────────────────────────────────────
    if block_type.startswith("heading"):
        level = int(block_type.replace("heading", ""))
        config = HEADING_STYLES.get(level, HEADING_STYLES[3])

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(config["space_before"])
        para.paragraph_format.space_after = Pt(6)

        run = para.add_run(text)
        run.bold = config["bold"]
        run.font.size = Pt(config["font_size"])
        run.font.color.rgb = config["color"]
        run.font.name = HEADING_FONT_NAME
        return

    # ── Regular paragraph ─────────────────────────────────────────
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.size = Pt(BODY_FONT_SIZE)
    run.font.name = BODY_FONT_NAME